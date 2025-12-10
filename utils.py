from docxtpl import DocxTemplate
from docx import Document
from io import BytesIO

def fill_template(template_file, data):
    """
    ממלא תבנית Word קיימת באמצעות ספריית docxtpl (Jinja2).
    שומר על העיצוב המקורי של הטופס בצורה מושלמת.
    """
    try:
        # טעינת התבנית מהקובץ שהועלה או מהנתיב המקומי
        doc = DocxTemplate(template_file)
        
        # הכנת הנתונים (Context)
        # הספרייה מצפה למילון פשוט. היא תחפש {{ key }} בתוך הוורד
        context = data
        
        # ביצוע הרינדור (החלפת המשתנים בערכים)
        doc.render(context)
        
        # שמירה לזיכרון (Buffer) כדי לא לשמור קובץ פיזי בשרת
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
        
    except Exception as e:
        # במקרה של תקלה (למשל התבנית לא תקינה), נחזיר הודעת שגיאה או מסמך ריק
        print(f"Error filling template: {e}")
        return None

def create_docx(data):
    """
    פונקציית גיבוי: יוצרת מסמך פשוט מאפס אם אין תבנית.
    שימושי למקרה שהמשתמש לא העלה תבנית ורוצה רק לראות את הנתונים.
    """
    doc = Document()
    doc.add_heading('סיכום תיק - נתונים גולמיים', 0)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ערך'
    hdr_cells[1].text = 'שדה'
    
    # הוספת הנתונים לטבלה
    for key, value in data.items():
        row_cells = table.add_row().cells
        # המרה למחרוזת וטיפול בעברית
        clean_key = str(key).replace("_", " ").title()
        val_str = str(value) if value is not None else ""
        
        row_cells[1].text = clean_key
        row_cells[0].text = val_str

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
